#!/usr/bin/env python
# coding: utf-8

# In[ ]:


from PyPDF2 import PdfReader, PdfWriter
import pdfplumber
from pdf2docx import Converter
import os
import docx

#读取文件夹内所有文件名
path = input('输入文件夹路径(结尾注意添加斜杠）：')
files= os.listdir(path)
#输入需要提取的章节：
a = input('请输入目标章节：')
Targetchapter = a + '\n'
b = input('请输入目标下一章节：')
Nextchapter = b + '\n'

#对所有文件进行遍历处理：
for filename in files:

    #初始化写入器与读取器：
    pdf_writer = PdfWriter()
    pdf_reader = PdfReader(path+filename)
    #进行关键词检索：
    with pdfplumber.open(path + filename) as pdf:
        print('正在读取中…')
        for i in range(len(pdf_reader.pages)):
            page = pdf.pages[i]

            if Targetchapter in page.extract_text():
                start = i
            elif Nextchapter in page.extract_text():
                end = i
        for i in range(start,end):
            pdf_writer.add_page(pdf_reader.pages[i])
    #提取出目标部分pdf内容：
    new = 'new'+filename
    with open(path+ new, 'wb') as out:
        pdf_writer.write(out)
    print('已完成目标内容提取，即将进入格式转换…')
    #将pdf内容转化为word形式：
    a = Converter(path+new) 
    a.convert(path+new.split('.')[0]+'.docx')
    a.close()
    print(filename + '完成转换，读取下一文件…')
print('已完成。')

print('继续提取小标题内容：')
TargetTitle = input('输入目标小标题：')
NextTitle = input('输入下一小标题：')
doc = docx.Document(path+new.split('.')[0]+'.docx')
print('正在定位目标位置……')
#获取目标小标题位置：
for i in range(0,len(doc.paragraphs)):
    if doc.paragraphs[i].text == TargetTitle:
        start = i
        print('find')
        print(start)
#获取目标下一小标题位置：
for i in range(0,len(doc.paragraphs)):
    if doc.paragraphs[i].text == NextTitle:
        end = i
        print('find')
        print(end)
#获取所有表格位置：
print(len(doc.tables))
#删除上方表格：
#遍历所有表格，比较是否在目标段落前：
for i in range(0,len(doc.tables)):
    # 获取段落和表格所在的_element对象
    paragraph_element = doc.paragraphs[start]._element
    table_element = doc.tables[0]._element
    # 获取段落和表格所在的_element对象在_document_part中的索引值
    paragraph_index_in_document_part = doc.part.element.body.index(paragraph_element)
    table_index_in_document_part = doc.part.element.body.index(table_element)
    # 根据索引值比较段落和表格的位置关系
    if paragraph_index_in_document_part < table_index_in_document_part:
        print("段落在表格前面")
#如果是，进行删除操作：      
    else:
        print("表格在段落前面")
        table = doc.tables[0]  
        table._element.getparent().remove(table._element)
#删除下方表格：
#遍历所有表格，比较是否在目标段落后：
for i in range(0,len(doc.tables)):
    # 获取段落和表格所在的_element对象
    paragraph_element = doc.paragraphs[end]._element
    table_element = doc.tables[0]._element
    # 获取段落和表格所在的_element对象在_document_part中的索引值
    paragraph_index_in_document_part = doc.part.element.body.index(paragraph_element)
    table_index_in_document_part = doc.part.element.body.index(table_element)
    # 根据索引值比较段落和表格的位置关系
    if paragraph_index_in_document_part < table_index_in_document_part:
    #如果是，进行删除操作：
        print("段落在表格前面")
        table = doc.tables[0]  
        table._element.getparent().remove(table._element)
    else:
        print("表格在段落前面")
        
#删除上方段落：
for i in range(0,start):
    paragraph = doc.paragraphs[0]  
    paragraph._element.getparent().remove(paragraph._element)
#删除下方段落：
for i in range(end - start-1,len(doc.paragraphs)):
    paragraph = doc.paragraphs[len(doc.paragraphs)-1]  
    paragraph._element.getparent().remove(paragraph._element)


    

doc.save('demo.docx')
print('finish')


# In[103]:


#截取具体小标题下辖内容
import docx
filename = input('输入文件名：')
TargetTitle = input('输入目标小标题：')
NextTitle = input('输入下一小标题：')
doc = docx.Document(filename)

#获取目标小标题位置：
for i in range(0,len(doc.paragraphs)):
    if doc.paragraphs[i].text == TargetTitle:
        start = i
        print('定位成功')
        global f
        f = 1
        break
    else:
        f = 2

if f != 1:
    print('定位失败，检查标题输入是否正确！')
#获取目标下一小标题位置：
for i in range(0,len(doc.paragraphs)):
    if doc.paragraphs[i].text == NextTitle:
        end = i
        print('截取成功')
        global b
        b = 1
        break
    else:
        b = 2
if b != 1:
    print('截取不成功,请检查标题是否输入正确！')
#获取所有表格位置：
#删除上方表格：
#遍历所有表格，比较是否在目标段落前：
for i in range(0,len(doc.tables)):
    # 获取段落和表格所在的_element对象
    paragraph_element = doc.paragraphs[start]._element
    table_element = doc.tables[0]._element
    # 获取段落和表格所在的_element对象在_document_part中的索引值
    paragraph_index_in_document_part = doc.part.element.body.index(paragraph_element)
    table_index_in_document_part = doc.part.element.body.index(table_element)
    # 根据索引值比较段落和表格的位置关系
    if paragraph_index_in_document_part < table_index_in_document_part:
        pass
    #:(段落在表格前面)
#如果是，进行删除操作：      
    else:
        #("表格在段落前面")
        table = doc.tables[0]  
        table._element.getparent().remove(table._element)
#删除下方表格：
#遍历所有表格，比较是否在目标段落后：
for i in range(0,len(doc.tables)):
    # 获取段落和表格所在的_element对象
    paragraph_element = doc.paragraphs[end]._element
    table_element = doc.tables[0]._element
    # 获取段落和表格所在的_element对象在_document_part中的索引值
    paragraph_index_in_document_part = doc.part.element.body.index(paragraph_element)
    table_index_in_document_part = doc.part.element.body.index(table_element)
    # 根据索引值比较段落和表格的位置关系
    if paragraph_index_in_document_part < table_index_in_document_part:
    #如果是，进行删除操作：
        #("段落在表格前面")
        table = doc.tables[0]  
        table._element.getparent().remove(table._element)
    else:
        pass
        #("表格在段落前面")

#删除上方段落：
for i in range(0,start):
    paragraph = doc.paragraphs[0]  
    paragraph._element.getparent().remove(paragraph._element)

#删除下方段落：
a = len(doc.paragraphs)
for i in range(end - start,len(doc.paragraphs)):
    a -= 1
    paragraph = doc.paragraphs[a]  
    paragraph._element.getparent().remove(paragraph._element)
   

    

doc.save('demo.docx')
print('finish')

   


# In[101]:


if paragraph.text == TargetTitle:
       start =doc.paragraphs.index(paragraph)
       print(start)


# In[ ]:




